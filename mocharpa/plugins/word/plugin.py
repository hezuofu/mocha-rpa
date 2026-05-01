"""Word plugin for the RPA framework.

Provides Word document read/write operations via ``python-docx``, integrated
with the framework's plugin lifecycle.

Usage::

    from mocharpa.plugins.word.plugin import WordPlugin
    from mocharpa.plugin.base import PluginManager

    mgr = PluginManager(context)
    word = WordPlugin()
    mgr.register(word)
    mgr.start_all()

    doc = word.open("template.docx")
    word.add_paragraph(doc, "Generated report")
    word.find_and_replace(doc, "{name}", "Alice")
    word.save("output.docx")
    mgr.shutdown_all()
"""

from __future__ import annotations

import logging
from typing import Any, Dict, List, Optional

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.table import Table as DocxTable

from mocharpa.plugin.base import Plugin

logger = logging.getLogger("rpa.word")


class WordPlugin:
    """Plugin for Word document operations using ``python-docx``.

    Attributes:
        name: Plugin identifier (``"word"``).
    """

    name = "word"

    def __init__(self) -> None:
        self._documents: Dict[str, Document] = {}
        self._context: Any = None

    # ------------------------------------------------------------------
    # Plugin lifecycle
    # ------------------------------------------------------------------

    def initialize(self, context: Any) -> None:
        self._context = context
        logger.info("WordPlugin initialized")

    def cleanup(self) -> None:
        for name in list(self._documents):
            self.close(name)
        logger.info("WordPlugin cleaned up")

    # ------------------------------------------------------------------
    # Document management
    # ------------------------------------------------------------------

    def open(self, path: str) -> Document:
        """Open an existing document.  Cached by filename stem.

        Returns:
            The :class:`docx.Document` instance.
        """
        name = self._key_from_path(path)
        if name in self._documents:
            return self._documents[name]
        doc = Document(path)
        self._documents[name] = doc
        logger.info("Opened document: %s", path)
        return doc

    def create(self, path: str = "document.docx") -> Document:
        """Create a new blank document."""
        name = self._key_from_path(path)
        doc = Document()
        self._documents[name] = doc
        logger.info("Created document: %s", path)
        return doc

    def save(self, name: str, path: Optional[str] = None) -> None:
        """Save a managed document.

        Args:
            name: Document key.
            path: Optional target path.
        """
        doc = self._documents.get(name)
        if doc is None:
            raise KeyError(f"Document '{name}' not found")
        target = path or f"{name}.docx"
        doc.save(target)
        logger.info("Saved document: %s", target)

    def close(self, name: str, *, save: bool = False) -> None:
        """Close a managed document."""
        doc = self._documents.pop(name, None)
        if doc is None:
            return
        if save:
            doc.save(name if "." in name else f"{name}.docx")
        logger.info("Closed document: %s", name)

    def get(self, name: str) -> Document:
        """Return a managed document by name."""
        doc = self._documents.get(name)
        if doc is None:
            raise KeyError(f"Document '{name}' not found")
        return doc

    def list_documents(self) -> List[str]:
        """Return names of all open documents."""
        return list(self._documents.keys())

    # ------------------------------------------------------------------
    # Paragraph operations
    # ------------------------------------------------------------------

    @staticmethod
    def add_paragraph(doc: Document, text: str, style: Optional[str] = None) -> Any:
        """Add a paragraph.  Returns the paragraph object."""
        return doc.add_paragraph(text, style=style)

    @staticmethod
    def add_heading(doc: Document, text: str, level: int = 1) -> Any:
        """Add a heading.  *level* is 1–9."""
        return doc.add_heading(text, level=level)

    @staticmethod
    def get_paragraphs(doc: Document) -> List[Any]:
        """Return all paragraphs in the document body."""
        return doc.paragraphs

    @staticmethod
    def get_text(doc: Document) -> str:
        """Return the full plain text of the document."""
        return "\n".join(p.text for p in doc.paragraphs)

    # ------------------------------------------------------------------
    # Table operations
    # ------------------------------------------------------------------

    @staticmethod
    def add_table(
        doc: Document,
        rows: int,
        cols: int,
        data: Optional[List[List[str]]] = None,
    ) -> DocxTable:
        """Add a table and optionally populate it with *data* (list of rows)."""
        table = doc.add_table(rows=rows, cols=cols, style="Table Grid")
        if data:
            for ri, row_data in enumerate(data):
                for ci, cell_value in enumerate(row_data):
                    if ri < rows and ci < cols:
                        table.cell(ri, ci).text = str(cell_value)
        return table

    @staticmethod
    def get_tables(doc: Document) -> List[DocxTable]:
        """Return all tables in the document."""
        return doc.tables

    @staticmethod
    def read_table(table: DocxTable) -> List[List[str]]:
        """Read table data as a list of rows."""
        return [[cell.text for cell in row.cells] for row in table.rows]

    # ------------------------------------------------------------------
    # Find & replace
    # ------------------------------------------------------------------

    @staticmethod
    def find_and_replace(doc: Document, old: str, new: str) -> int:
        """Replace all occurrences of *old* with *new* in paragraphs and tables.

        Returns:
            Number of replacements made.
        """
        count = 0

        # Paragraphs
        for para in doc.paragraphs:
            if old in para.text:
                # python-docx inline runs make this non-trivial;
                # use a simple text-level approach for the main runs
                for run in para.runs:
                    if old in run.text:
                        run.text = run.text.replace(old, new)
                        count += 1

        # Tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        for run in para.runs:
                            if old in run.text:
                                run.text = run.text.replace(old, new)
                                count += 1

        logger.info("find_and_replace: replaced %d occurrences of '%s'", count, old)
        return count

    # ------------------------------------------------------------------
    # Image
    # ------------------------------------------------------------------

    @staticmethod
    def add_picture(doc: Document, image_path: str, width_inches: float = 3.0) -> Any:
        """Add an image at the end of the document."""
        return doc.add_picture(image_path, width=Inches(width_inches))

    # ------------------------------------------------------------------
    # Helpers
    # ------------------------------------------------------------------

    @staticmethod
    def _key_from_path(path: str) -> str:
        import os
        return os.path.splitext(os.path.basename(path))[0]
